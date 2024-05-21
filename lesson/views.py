from django.shortcuts import render
from .models import Lesson
from django.http import HttpResponseRedirect
from django.core.files.storage import FileSystemStorage
from django.urls import reverse
from django.http import HttpResponse, HttpResponseRedirect
from django.urls import reverse
from django.core.files.storage import FileSystemStorage
from .models import Lesson
import requests
import json
from http import HTTPStatus
import dashscope
import re
import os
dashscope.api_key="sk-ed0c4bfbafb541a0812bda5b311cebdf"
def get_access_token():
    """
    使用 AK，SK 生成鉴权签名（Access Token）
    :return: access_token，或是None(如果错误)
    """
    url = "https://aip.baidubce.com/oauth/2.0/token"
    params = {"grant_type": "client_credentials", "client_id": API_KEY, "client_secret": SECRET_KEY}
    return str(requests.post(url, params=params).json().get("access_token"))
def lesson_list(request):
    lessons = Lesson.objects.all()
    return render(request, 'lesson_list.html', {'lessons': lessons})


from django.http import HttpResponse
from django.core.files.storage import FileSystemStorage
from .models import Lesson  # 假设你的模型是这样导入的
from docx import Document#python-docx
from shiyan.models import Shiyan
import difflib
def upload_file(request, id):
    if request.method == 'POST':
        uid = request.session.get('user_id')

        if 'file' in request.FILES:
            file = request.FILES['file']
            if file.name.endswith('.docx'):
                fs = FileSystemStorage()
                filename = fs.save(file.name, file)
                file_path = fs.path(filename)

                doc = Document(file_path)
                content = []
                for paragraph in doc.paragraphs:
                    content.append(paragraph.text)

                # 读取表格
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            content.append(cell.text)
                content = ''.join(content)

                # 计算内容重复度
                existing_texts = Shiyan.objects.filter(lid=id).values_list('text', flat=True)
                similarity_scores = [difflib.SequenceMatcher(None, content, old_text).ratio() for old_text in existing_texts]
                average_similarity = max(similarity_scores) if similarity_scores else 0
                print(average_similarity)
                # 处理文件内容
                chuli_result, score = chuli(content)  # 假设chuli函数处理内容并返回结果和分数
                ff=os.path.basename(file_path)
                # 检查是否已有相同lid和uid的记录
                existing_record = Shiyan.objects.filter(lid=id, uid=uid).first()
                if existing_record:
                    # 更新现有记录
                    existing_record.pingyu = chuli_result
                    existing_record.score = score
                    existing_record.file = fs.url(filename)
                    existing_record.text = content
                    existing_record.chong = average_similarity
                    existing_record.url=ff
                    existing_record.save()
                    return HttpResponse("文件已更新，处理结果: {chuli_result}")
                else:
                    # 创建新记录
                    new_record = Shiyan(
                        url=ff,
                        lid=id,
                        uid=uid,
                        pingyu=chuli_result,
                        score=score,
                        file=fs.url(filename),
                        text=content,
                        chong=average_similarity
                    )
                    new_record.save()
                    return HttpResponse("文件已处理并保存，处理结果: {chuli_result}")
            else:
                return HttpResponse("上传失败：只接受 .docx 文件。")
        else:
            return HttpResponse("上传失败：没有找到文件。")
    else:
        return HttpResponse("错误请求：只支持 POST 方法。")
# def upload_file(request, id):
#     # 确保是 POST 方法
#     if request.method == 'POST':
#         uid=request.session.get('user_id')
#         # 检查是否有文件被上传
#         if 'file' in request.FILES:
#             file = request.FILES['file']
#             # 检查文件类型
#             if file.name.endswith('.docx'):
#                 fs = FileSystemStorage()
#                 filename = fs.save(file.name, file)
#                 file_path = fs.path(filename)  # 获取文件的系统路径
#
#                 # 使用python-docx读取Word文件内容
#                 doc = Document(file_path)
#                 # 读取段落
#                 content = []
#                 for paragraph in doc.paragraphs:
#                     content.append(paragraph.text)
#
#                 # 读取表格
#                 for table in doc.tables:
#                     for row in table.rows:
#                         for cell in row.cells:
#                             content.append(cell.text)
#                 content = ''.join(content)
#                 # 调用处理函数处理文件内容
#                 chuli_result, score = chuli(content)  # 假设 chuli 函数现在接收文件内容并返回处理结果和分数
#
#                 # 创建并保存模型实例（这里我们仍然保存文件的URL，但你也可以选择保存文件内容或其他信息）
#                 lesson_file = Shiyan(
#                     lid=id,
#                     uid=uid,
#                     pingyu=chuli_result,
#                     score=score,
#                     file=fs.url(filename) , # 或者你可以选择保存其他与文件相关的信息
#                     text=content
#                 )
#                 lesson_file.save()
#
#                 return HttpResponse(f"文件已处理，处理结果: {chuli_result}")
#             else:
#                 # 返回错误信息
#                 return HttpResponse("上传失败：只接受 .docx 文件。")
#         else:
#             return HttpResponse("上传失败：没有找到文件。")
#     else:
#         return HttpResponse("错误请求：只支持 POST 方法。")
#

def chuli(content):

    response = dashscope.Generation.call(
        model=dashscope.Generation.Models.qwen_turbo,
        prompt='你现在是一名大学教授，请给你下面的一份C语言编程课实验报告满分100分，请从下面我给出的角度进行分析：'
               '1.如果字数小于300 请给低分。'
               '2.字数小于100且无有效内容给59分' 
               '3.如果有3张图或以上请给高一点'
               '4.如果与C语言内容无关请给低分'
               '如果无法获取相关内容就给0分'
               '做出评论，并且打分，满分100分，打分要严格些严格些!!!并且精确到个位!!!你给我的最后一句话一定必须是这样的:因此，总分为（填分数）分}'+content
    )
    # print(response.output)  # The output text
    # print(response.usage)  # The usage information
    score_match = re.search(r'总分为(\d+)分', response.output.text)

    # 检查是否匹配到分数

    if score_match:

        # 提取并打印总分

        total_score = score_match.group(1)

        print(f"提取的总分为：{total_score}分")

    else:
        total_score='暂未获取'
        print("未找到总分信息。")

    return response.output.text, total_score  # 返回评语和分数
def view_scores(request, id):
    # 假设我们有一个模型 Score 来存储得分，与 Lesson 模型有关联
    # from .models import Score
    # scores = Score.objects.filter(lesson_id=id)

    # 为简单起见，这里我们用一个静态的数据示例
    experiments = Shiyan.objects.filter(lid=id,uid=request.session.get('user_id')).values('lid', 'uid', 'score', 'pingyu', 'chong', 'url')
    print(experiments)
    # 构建一个列表，其中包含所需的每个实验的信息
    scores = [
        {'student': experiment['uid'], 'score': experiment['score'], 'comment': experiment['pingyu'], 'chong': str(float(experiment['chong'])*100), 'url': experiment['url']}
        for experiment in experiments
    ]

    # 渲染HTML模板，传递scores列表


    return render(request, 'scores.html', {'scores': scores, 'lesson_id': id})

